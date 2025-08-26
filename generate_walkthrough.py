#!/usr/bin/env python3

import os
import re
from typing import Dict, List, Tuple

try:
	from docx import Document
	from docx.shared import Pt
	from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
	print("[WARN] python-docx not available yet: {}".format(e))
	Document = None

ROOT_DIR = "/workspace/Auto-Insurance"
OUTPUT_DOCX = "/workspace/ASP.NET_Core_Project_Walkthrough.docx"

EXCLUDE_DIR_NAMES = {".git", "bin", "obj", ".vs", "node_modules", ".venv", "packages"}
INCLUDE_EXTENSIONS = {".cs", ".cshtml", ".csproj", ".json", ".razor", ".config"}

# --------------------------- File Discovery ---------------------------

def discover_files(root_dir: str) -> List[str]:
	matching_files: List[str] = []
	for dirpath, dirnames, filenames in os.walk(root_dir):
		# prune excluded directories in-place
		dirnames[:] = [d for d in dirnames if d not in EXCLUDE_DIR_NAMES]
		for filename in filenames:
			_, ext = os.path.splitext(filename)
			if ext.lower() in INCLUDE_EXTENSIONS:
				matching_files.append(os.path.join(dirpath, filename))
	return sorted(matching_files)

# --------------------------- Explanation Heuristics ---------------------------

def detect_language_and_comment(file_path: str) -> Tuple[str, str]:
	_, ext = os.path.splitext(file_path.lower())
	if ext in {".cs", ".cshtml", ".razor"}:
		return ("csharp_or_razor", " // ")
	elif ext in {".json", ".config"}:
		return ("json_or_config", " // ")
	elif ext in {".csproj"}:
		return ("xml", " // ")
	else:
		return ("plain", " // ")

# Patterns compiled once for speed
USING_RE = re.compile(r"^\s*using\s+([\w\.\=]+)\s*;\s*$")
NAMESPACE_RE = re.compile(r"^\s*namespace\s+([\w\.]+)\s*\{?\s*$")
CLASS_RE = re.compile(r"^\s*(public|internal|protected|private)?\s*(abstract\s+|static\s+|sealed\s+)?\s*class\s+([A-Za-z_][A-Za-z0-9_]*)\s*(?::\s*([^{]+))?\s*\{?\s*$")
INTERFACE_RE = re.compile(r"^\s*public\s+interface\s+([A-Za-z_][A-Za-z0-9_]*)")
METHOD_RE = re.compile(r"^\s*(public|private|protected|internal)\s+(async\s+)?([\w<>\[\]]+)\s+([A-Za-z_][A-Za-z0-9_]*)\s*\(([^)]*)\)")
ATTR_RE = re.compile(r"^\s*\[([A-Za-z_][A-Za-z0-9_]*(?:\([^\]]*\))?)\]\s*$")
ROUTE_ATTR_RE = re.compile(r"\[Route\((.*?)\)\]")
HTTP_ATTR_RE = re.compile(r"\[(HttpGet|HttpPost|HttpPut|HttpDelete|HttpPatch)(?:\((.*?)\))?\]")
RETURN_VIEW_RE = re.compile(r"return\s+View\((.*?)\)\s*;")
RETURN_JSON_RE = re.compile(r"return\s+Json\((.*?)\)\s*;")
NEW_DBCTX_RE = re.compile(r"DbContext|AutoInsuranceDbContext|AutoInsuranceApiDbContext")
SERVICE_CALL_RE = re.compile(r"(\w+Service)\.")

RAZOR_MODEL_RE = re.compile(r"^\s*@model\s+(.*)")
RAZOR_USING_RE = re.compile(r"^\s*@using\s+(.*)")
RAZOR_SECTION_RE = re.compile(r"^\s*@section\s+(\w+)\s*\{")
RAZOR_IF_RE = re.compile(r"^\s*@if\s*\((.*)\)\s*\{")
RAZOR_FOR_RE = re.compile(r"^\s*@for(each)?\s*\((.*)\)\s*\{")

JSON_KEY_RE = re.compile(r"^\s*\"([A-Za-z0-9_\-]+)\"\s*:\s*(.*)")
XML_TAG_RE = re.compile(r"^\s*<([A-Za-z0-9_\-\.]+)(\s+[^>]*)?>")
PACKAGE_REF_RE = re.compile(r"<PackageReference\s+Include=\"([^\"]+)\"\s+Version=\"([^\"]+)\"\s*/>")


def explain_csharp_or_razor_line(file_path: str, line: str) -> str:
	line_stripped = line.strip()
	if not line_stripped:
		return "Blank line for readability"

	m = USING_RE.match(line)
	if m:
		return f"Imports the {m.group(1)} namespace so its types are available"

	m = NAMESPACE_RE.match(line)
	if m:
		return f"Declares the namespace {m.group(1)} to logically group related code"

	m = ATTR_RE.match(line)
	if m:
		attr = m.group(1)
		# Common attribute hints
		if attr.startswith("ApiController"):
			return "Marks this class as a Web API controller for automatic behaviors"
		if attr.startswith("Authorize"):
			return "Requires the user to be authorized to access the decorated action/class"
		if HTTP_ATTR_RE.search(line):
			http_m = HTTP_ATTR_RE.search(line)
			verb = http_m.group(1)
			path = http_m.group(2) or ""
			return f"Maps this action to HTTP {verb} {path}".strip()
		if ROUTE_ATTR_RE.search(line):
			route = ROUTE_ATTR_RE.search(line).group(1)
			return f"Defines a route template {route} for this controller/action"
		return f"Applies the {attr} attribute which modifies behavior/metadata"

	m = CLASS_RE.match(line)
	if m:
		cls_name = m.group(3)
		inherit = (m.group(4) or "").strip()
		if inherit:
			return f"Declares class {cls_name} inheriting from {inherit}"
		return f"Declares class {cls_name}"

	m = INTERFACE_RE.match(line)
	if m:
		return f"Declares interface {m.group(1)} (contract of methods/properties)"

	m = METHOD_RE.match(line)
	if m:
		access = m.group(1)
		is_async = bool(m.group(2))
		ret = m.group(3)
		name = m.group(4)
		params_sig = m.group(5)
		async_note = " asynchronously" if is_async else ""
		return f"{access.capitalize()} method {name} returns {ret}{async_note} with parameters ({params_sig})"

	if RETURN_VIEW_RE.search(line):
		return "Returns a Razor View to render HTML to the client"
	if RETURN_JSON_RE.search(line):
		return "Returns JSON data as the HTTP response"
	if NEW_DBCTX_RE.search(line):
		return "Uses the EF Core DbContext to query or modify the database"
	if SERVICE_CALL_RE.search(line):
		svc = SERVICE_CALL_RE.search(line).group(1)
		return f"Calls {svc} to execute business logic"

	if line_stripped.startswith("@model"):
		m = RAZOR_MODEL_RE.match(line)
		return f"Specifies the model type for this view: {m.group(1) if m else ''}"
	if line_stripped.startswith("@using"):
		m = RAZOR_USING_RE.match(line)
		return f"Imports namespace in the view: {m.group(1) if m else ''}"
	if line_stripped.startswith("@{"):
		return "Begins a Razor code block for server-side logic"
	if RAZOR_SECTION_RE.match(line):
		name = RAZOR_SECTION_RE.match(line).group(1)
		return f"Defines a Razor section named {name} to inject content into layout"
	if RAZOR_IF_RE.match(line):
		return "Razor conditional block that runs only when the condition is true"
	if RAZOR_FOR_RE.match(line):
		return "Razor loop to iterate over a collection and render repeated markup"

	if line_stripped.startswith("<") and line_stripped.endswith(">"):
		return "HTML/Razor markup element rendered to the page"

	return "Code statement that contributes to the application's behavior"


def explain_json_or_config_line(line: str) -> str:
	line_stripped = line.strip()
	if not line_stripped:
		return "Blank line for readability"
	m = JSON_KEY_RE.match(line)
	if m:
		key = m.group(1)
		return f"Configures the setting '{key}'"
	if line_stripped.startswith("{"):
		return "Begins a JSON object"
	if line_stripped.startswith("}"):
		return "Ends a JSON object"
	if line_stripped.startswith("["):
		return "Begins a JSON array"
	if line_stripped.startswith("]"):
		return "Ends a JSON array"
	return "JSON/config line setting application behavior"


def explain_xml_line(line: str) -> str:
	line_stripped = line.strip()
	if not line_stripped:
		return "Blank line for readability"
	m = PACKAGE_REF_RE.search(line)
	if m:
		return f"Adds NuGet package dependency '{m.group(1)}' version {m.group(2)}"
	m = XML_TAG_RE.match(line)
	if m:
		tag = m.group(1)
		return f"XML element <{tag}> configuring project/build settings"
	if line_stripped.startswith("<!--"):
		return "XML comment"
	return "XML/project configuration line"


def explain_line(file_path: str, line: str) -> str:
	lang, _ = detect_language_and_comment(file_path)
	if lang == "csharp_or_razor":
		return explain_csharp_or_razor_line(file_path, line)
	elif lang == "json_or_config":
		return explain_json_or_config_line(line)
	elif lang == "xml":
		return explain_xml_line(line)
	else:
		return "Line explanation"

# --------------------------- Document Builder ---------------------------

def add_heading(document: Document, text: str, level: int) -> None:
	p = document.add_heading(text, level=level)
	p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_code_line_with_comment(document: Document, file_path: str, line: str) -> None:
	lang, comment_delim = detect_language_and_comment(file_path)
	# Keep original indentation; append inline comment explanation
	explanation = explain_line(file_path, line)
	paragraph = document.add_paragraph()
	run_code = paragraph.add_run(line.rstrip("\n"))
	run_code.font.name = "Consolas"
	run_code.font.size = Pt(10)
	# Ensure there is at least one space before the inline comment
	separator = ""
	if not line.rstrip("\n").endswith(" "):
		separator = " "
	run_comment = paragraph.add_run(f"{separator}{comment_delim}{explanation}")
	run_comment.font.name = "Consolas"
	run_comment.font.size = Pt(10)


def add_file_section(document: Document, file_path: str) -> None:
	relative = os.path.relpath(file_path, ROOT_DIR)
	add_heading(document, f"File: {relative}", level=2)
	try:
		with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
			for line in f:
				add_code_line_with_comment(document, file_path, line)
	except Exception as e:
		paragraph = document.add_paragraph()
		run = paragraph.add_run(f"[Could not read file due to error: {e}]")
		run.font.name = "Consolas"
		run.font.size = Pt(10)


def gather_structure(files: List[str]) -> Dict[str, List[str]]:
	"""Group files by top-level project directory for organization."""
	structure: Dict[str, List[str]] = {}
	for path in files:
		# project is the immediate child of ROOT_DIR (e.g., Auto_Insurance_System, AutoInsuranceSystemAPI)
		rel = os.path.relpath(path, ROOT_DIR)
		project_dir = rel.split(os.sep)[0] if os.sep in rel else rel
		structure.setdefault(project_dir, []).append(path)
	# sort the file lists for stable ordering
	for key in structure:
		structure[key] = sorted(structure[key])
	return dict(sorted(structure.items(), key=lambda kv: kv[0]))


def build_flow_summary(document: Document, files: List[str]) -> None:
	add_heading(document, "Application Working Flow Summary", level=1)
	p = document.add_paragraph()
	r = p.add_run(
		"This solution contains an ASP.NET Core MVC web app (Auto_Insurance_System) and a Web API (AutoInsuranceSystemAPI). "
		"The overall flow follows the standard MVC and API patterns. Below is a step-by-step overview."
	)
	r.font.size = Pt(11)

	steps: List[str] = [
		"A user navigates to a URL or triggers an HTTP request.",
		"Routing maps the request to a controller and action method based on route templates and HTTP verbs.",
		"The controller coordinates work: it validates input, calls services for business logic, and interacts with the database via EF Core DbContext (e.g., AutoInsuranceDbContext, AutoInsuranceApiDbContext).",
		"Services encapsulate business rules and data operations; they use DbContext to query and update entities like Policy, Claim, Payment, SupportTicket, and Users.",
		"Models (POCOs) represent the domain data and are tracked by EF Core; view models or DTOs may be used to shape data returned to views or APIs.",
		"For MVC, the controller returns a View along with a model. Razor views (.cshtml) render HTML using the model data.",
		"For Web API, the controller returns an IActionResult or a typed result (e.g., IEnumerable<Policy>), which is serialized to JSON.",
		"Filters (e.g., authorization and exception filters) run before/after actions to enforce security and handle cross-cutting concerns.",
		"Finally, the response (HTML or JSON) is sent back to the client." 
	]
	for idx, step in enumerate(steps, start=1):
		bullet = document.add_paragraph(style=None)
		run = bullet.add_run(f"{idx}. {step}")
		run.font.size = Pt(11)

	add_heading(document, "How parts connect", level=2)
	connections: List[str] = [
		"Controllers depend on Services via constructor injection (DI).",
		"Services depend on the EF Core DbContext to query/update the database.",
		"Models define the shape of tables and are used by DbContext and passed to Views.",
		"Views (.cshtml) render the UI by reading properties from the model.",
		"In the API project, controllers return data as JSON; in the MVC project, controllers return Views.",
	]
	for c in connections:
		para = document.add_paragraph()
		run = para.add_run(f"- {c}")
		run.font.size = Pt(11)


def generate_document() -> None:
	if Document is None:
		raise RuntimeError("python-docx is not installed. Please install and rerun.")

	document = Document()
	add_heading(document, "ASP.NET Core Project Walkthrough", level=0)
	intro = document.add_paragraph()
	run = intro.add_run(
		"This document explains each file with simple inline comments and concludes with a working flow summary."
	)
	run.font.size = Pt(11)

	files = discover_files(ROOT_DIR)
	structure = gather_structure(files)

	for project, project_files in structure.items():
		add_heading(document, f"Project: {project}", level=1)
		# group by subfolder (Controllers, Models, Views, Services, Data, Filters, etc.)
		folders: Dict[str, List[str]] = {}
		for f in project_files:
			rel = os.path.relpath(f, os.path.join(ROOT_DIR, project))
			folder = rel.split(os.sep)[0] if os.sep in rel else "."
			folders.setdefault(folder, []).append(f)
		for folder_name in sorted(folders.keys()):
			add_heading(document, f"Folder: {folder_name}", level=2)
			for file_path in sorted(folders[folder_name]):
				add_file_section(document, file_path)

	# Flow summary
	build_flow_summary(document, files)

	document.save(OUTPUT_DOCX)
	print(f"[OK] Walkthrough document generated at: {OUTPUT_DOCX}")


if __name__ == "__main__":
	generate_document()